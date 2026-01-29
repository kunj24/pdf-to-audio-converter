"""
Document Converter Module
Supports multiple input formats: PDF, DOCX, TXT, EPUB, HTML, Images (OCR)
"""

import os
import re
import logging
from typing import Optional, Tuple, List, Dict
from dataclasses import dataclass
from enum import Enum
import tempfile

logger = logging.getLogger(__name__)

# Import PDF reader
try:
    from PyPDF2 import PdfReader
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False
    logger.warning("PyPDF2 not available")

# Import docx reader
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available - DOCX support disabled")

# Import OCR
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False
    logger.warning("pytesseract/PIL not available - OCR support disabled")

# Import PDF to image converter
try:
    import pdf2image
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False
    logger.warning("pdf2image not available - scanned PDF OCR disabled")

# Import epub reader
try:
    import ebooklib
    from ebooklib import epub
    from bs4 import BeautifulSoup
    EPUB_AVAILABLE = True
except ImportError:
    EPUB_AVAILABLE = False
    logger.warning("ebooklib/bs4 not available - EPUB support disabled")

# Import HTML parser
try:
    from bs4 import BeautifulSoup
    BS4_AVAILABLE = True
except ImportError:
    BS4_AVAILABLE = False


class DocumentFormat(Enum):
    """Supported document formats"""
    PDF = 'pdf'
    DOCX = 'docx'
    DOC = 'doc'
    TXT = 'txt'
    EPUB = 'epub'
    HTML = 'html'
    HTM = 'htm'
    IMAGE = 'image'  # JPG, PNG, etc. (OCR)
    URL = 'url'


@dataclass
class DocumentInfo:
    """Document metadata"""
    format: DocumentFormat
    page_count: int
    word_count: int
    has_images: bool
    has_tables: bool
    language: Optional[str]
    title: Optional[str]
    author: Optional[str]
    is_scanned: bool = False


class DocumentConverter:
    """Convert various document formats to text"""
    
    def __init__(self):
        self.supported_formats = self._get_supported_formats()
    
    def _get_supported_formats(self) -> Dict[str, bool]:
        """Get list of supported formats based on available libraries"""
        return {
            'pdf': PYPDF2_AVAILABLE,
            'doc': DOCX_AVAILABLE,  # python-docx can handle some .doc files
            'docx': DOCX_AVAILABLE,
            'txt': True,
            'epub': EPUB_AVAILABLE,
            'html': BS4_AVAILABLE,
            'htm': BS4_AVAILABLE,
            'image_ocr': OCR_AVAILABLE,
            'scanned_pdf': OCR_AVAILABLE and PDF2IMAGE_AVAILABLE,
        }
    
    def get_format(self, file_path: str) -> DocumentFormat:
        """Detect document format from file extension (robust handling)"""
        # Get extension and clean it thoroughly
        ext = os.path.splitext(file_path)[1].lower().lstrip('.').strip()
        
        # Remove any non-alphanumeric characters
        ext = ''.join(c for c in ext if c.isalnum())
        
        format_map = {
            'pdf': DocumentFormat.PDF,
            'docx': DocumentFormat.DOCX,
            'doc': DocumentFormat.DOC,
            'txt': DocumentFormat.TXT,
            'text': DocumentFormat.TXT,
            'epub': DocumentFormat.EPUB,
            'html': DocumentFormat.HTML,
            'htm': DocumentFormat.HTM,
            'jpg': DocumentFormat.IMAGE,
            'jpeg': DocumentFormat.IMAGE,
            'png': DocumentFormat.IMAGE,
            'bmp': DocumentFormat.IMAGE,
            'tiff': DocumentFormat.IMAGE,
            'tif': DocumentFormat.IMAGE,
            'gif': DocumentFormat.IMAGE,
        }
        
        detected_format = format_map.get(ext)
        
        if not detected_format:
            # Try to detect from file content if extension is unclear
            if ext in ['', 'tmp', 'download']:
                try:
                    # Read first few bytes to detect PDF signature
                    with open(file_path, 'rb') as f:
                        header = f.read(5)
                        if header.startswith(b'%PDF'):
                            return DocumentFormat.PDF
                except:
                    pass
            return DocumentFormat.TXT  # Default to text
        
        return detected_format
    
    def convert(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None,
        use_ocr: bool = False
    ) -> Tuple[str, DocumentInfo]:
        """
        Convert document to text
        
        Args:
            file_path: Path to document
            start_page: Start page (1-indexed, for PDFs)
            end_page: End page (1-indexed, for PDFs)
            use_ocr: Force OCR for PDF (useful for scanned documents)
        
        Returns:
            Tuple of (extracted_text, document_info)
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        doc_format = self.get_format(file_path)
        
        converters = {
            DocumentFormat.PDF: self._convert_pdf,
            DocumentFormat.DOCX: self._convert_docx,
            DocumentFormat.DOC: self._convert_docx,  # Try using DOCX handler for DOC
            DocumentFormat.TXT: self._convert_txt,
            DocumentFormat.EPUB: self._convert_epub,
            DocumentFormat.HTML: self._convert_html,
            DocumentFormat.HTM: self._convert_html,
            DocumentFormat.IMAGE: self._convert_image_ocr,
        }
        
        converter = converters.get(doc_format)
        if not converter:
            raise ValueError(f"Unsupported format: {doc_format}")
        
        # Special handling for PDF with OCR option
        if doc_format == DocumentFormat.PDF and use_ocr:
            return self._convert_pdf_ocr(file_path, start_page, end_page)
        
        return converter(file_path, start_page, end_page)
    
    def _convert_pdf(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert PDF to text with optimized extraction"""
        if not PYPDF2_AVAILABLE:
            raise RuntimeError("PyPDF2 is required for PDF conversion")
        
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        
        # Handle page range
        start = (start_page - 1) if start_page else 0
        end = end_page if end_page else total_pages
        start = max(0, start)
        end = min(total_pages, end)
        
        # Optimized: Extract text with batch processing for large docs
        text_chunks = []
        has_images = False
        
        # Use batch processing for better memory efficiency
        BATCH_SIZE = 50  # Process 50 pages at a time
        for batch_start in range(start, end, BATCH_SIZE):
            batch_end = min(batch_start + BATCH_SIZE, end)
            
            for i in range(batch_start, batch_end):
                page = reader.pages[i]
                # Optimized: Extract text and skip empty pages
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text_chunks.append(page_text.strip())
                
                # Check for images (simple check without iterating)
                if not has_images:
                    try:
                        # Safer way: just check if resources exist
                        if hasattr(page, 'images') and len(page.images) > 0:
                            has_images = True
                    except:
                        pass  # Skip image detection on error
        
        # Optimized: Join with better spacing
        full_text = "\n\n".join(text_chunks)
        
        # Check if document appears to be scanned (very little text per page)
        avg_chars_per_page = len(full_text) / max(1, end - start)
        is_scanned = avg_chars_per_page < 100 and has_images
        
        # Extract metadata
        metadata = reader.metadata or {}
        
        info = DocumentInfo(
            format=DocumentFormat.PDF,
            page_count=total_pages,
            word_count=len(full_text.split()),
            has_images=has_images,
            has_tables='\t' in full_text,
            language=None,
            title=metadata.get('/Title'),
            author=metadata.get('/Author'),
            is_scanned=is_scanned
        )
        
        return full_text, info
    
    def _convert_pdf_ocr(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert scanned PDF using OCR"""
        if not OCR_AVAILABLE or not PDF2IMAGE_AVAILABLE:
            raise RuntimeError("pytesseract and pdf2image required for OCR")
        
        # Convert PDF pages to images
        first_page = start_page or 1
        last_page = end_page
        
        images = pdf2image.convert_from_path(
            file_path,
            first_page=first_page,
            last_page=last_page,
            dpi=300  # Higher DPI for better OCR
        )
        
        # OCR each page
        text_chunks = []
        for i, image in enumerate(images):
            page_text = pytesseract.image_to_string(image)
            text_chunks.append(page_text)
        
        full_text = "\n\n".join(text_chunks)
        
        # Get total page count from original PDF
        reader = PdfReader(file_path)
        total_pages = len(reader.pages)
        
        info = DocumentInfo(
            format=DocumentFormat.PDF,
            page_count=total_pages,
            word_count=len(full_text.split()),
            has_images=True,
            has_tables=False,
            language=None,
            title=None,
            author=None,
            is_scanned=True
        )
        
        return full_text, info
    
    def _convert_docx(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert DOCX/DOC to text"""
        if not DOCX_AVAILABLE:
            raise RuntimeError("python-docx required for DOCX/DOC conversion")
        
        try:
            doc = DocxDocument(file_path)
        except Exception as e:
            # If it's a .doc file, python-docx might fail
            ext = os.path.splitext(file_path)[1].lower()
            if ext == '.doc':
                raise RuntimeError(f"Legacy .doc format not fully supported. Please convert to .docx first. Error: {e}")
            raise
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        has_tables = len(doc.tables) > 0
        for table in doc.tables:
            for row in table.rows:
                row_text = '\t'.join(cell.text for cell in row.cells)
                if row_text.strip():
                    paragraphs.append(row_text)
        
        full_text = "\n\n".join(paragraphs)
        
        # Extract metadata
        core_props = doc.core_properties
        
        # Determine format from file extension
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        doc_format = DocumentFormat.DOC if ext == 'doc' else DocumentFormat.DOCX
        
        info = DocumentInfo(
            format=doc_format,
            page_count=1,  # DOCX doesn't have fixed pages
            word_count=len(full_text.split()),
            has_images=len(doc.inline_shapes) > 0,
            has_tables=has_tables,
            language=None,
            title=core_props.title,
            author=core_props.author,
            is_scanned=False
        )
        
        return full_text, info
    
    def _convert_txt(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert TXT to text"""
        # Try different encodings
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        text = None
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                break
            except (UnicodeDecodeError, UnicodeError):
                continue
        
        if text is None:
            raise RuntimeError("Could not decode text file")
        
        info = DocumentInfo(
            format=DocumentFormat.TXT,
            page_count=1,
            word_count=len(text.split()),
            has_images=False,
            has_tables='\t' in text,
            language=None,
            title=os.path.basename(file_path),
            author=None,
            is_scanned=False
        )
        
        return text, info
    
    def _convert_epub(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert EPUB to text"""
        if not EPUB_AVAILABLE:
            raise RuntimeError("ebooklib and beautifulsoup4 required for EPUB")
        
        book = epub.read_epub(file_path)
        
        text_chunks = []
        chapter_count = 0
        
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                
                # Remove scripts and styles
                for script in soup(['script', 'style']):
                    script.decompose()
                
                text = soup.get_text(separator='\n')
                if text.strip():
                    text_chunks.append(text)
                    chapter_count += 1
        
        full_text = "\n\n".join(text_chunks)
        
        # Extract metadata
        title = book.get_metadata('DC', 'title')
        author = book.get_metadata('DC', 'creator')
        
        info = DocumentInfo(
            format=DocumentFormat.EPUB,
            page_count=chapter_count,
            word_count=len(full_text.split()),
            has_images=False,
            has_tables=False,
            language=None,
            title=title[0][0] if title else None,
            author=author[0][0] if author else None,
            is_scanned=False
        )
        
        return full_text, info
    
    def _convert_html(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert HTML to text"""
        if not BS4_AVAILABLE:
            raise RuntimeError("beautifulsoup4 required for HTML conversion")
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove scripts, styles, and other non-content elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
            element.decompose()
        
        # Get text
        text = soup.get_text(separator='\n')
        
        # Clean up whitespace
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        # Extract title
        title_tag = soup.find('title')
        title = title_tag.string if title_tag else None
        
        info = DocumentInfo(
            format=DocumentFormat.HTML,
            page_count=1,
            word_count=len(text.split()),
            has_images=len(soup.find_all('img')) > 0,
            has_tables=len(soup.find_all('table')) > 0,
            language=soup.html.get('lang') if soup.html else None,
            title=title,
            author=None,
            is_scanned=False
        )
        
        return text, info
    
    def _convert_image_ocr(
        self,
        file_path: str,
        start_page: Optional[int] = None,
        end_page: Optional[int] = None
    ) -> Tuple[str, DocumentInfo]:
        """Convert image to text using OCR"""
        if not OCR_AVAILABLE:
            raise RuntimeError("pytesseract and Pillow required for image OCR")
        
        # Open and preprocess image
        image = Image.open(file_path)
        
        # Convert to RGB if necessary
        if image.mode != 'RGB':
            image = image.convert('RGB')
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        
        info = DocumentInfo(
            format=DocumentFormat.IMAGE,
            page_count=1,
            word_count=len(text.split()),
            has_images=True,
            has_tables=False,
            language=None,
            title=os.path.basename(file_path),
            author=None,
            is_scanned=True
        )
        
        return text, info
    
    def detect_if_scanned(self, file_path: str) -> bool:
        """Detect if a PDF is scanned (image-based)"""
        if not PYPDF2_AVAILABLE:
            return False
        
        try:
            reader = PdfReader(file_path)
            
            # Check first few pages
            pages_to_check = min(3, len(reader.pages))
            total_text = 0
            
            for i in range(pages_to_check):
                page_text = reader.pages[i].extract_text() or ""
                total_text += len(page_text)
            
            # If average text per page is very low, likely scanned
            avg_chars = total_text / pages_to_check
            return avg_chars < 100
            
        except Exception as e:
            logger.warning(f"Could not detect if PDF is scanned: {e}")
            return False


class URLConverter:
    """Convert web URLs to text"""
    
    def __init__(self):
        self.session = None
    
    def convert(self, url: str) -> Tuple[str, DocumentInfo]:
        """Fetch and convert URL content to text"""
        try:
            import requests
        except ImportError:
            raise RuntimeError("requests library required for URL conversion")
        
        if not BS4_AVAILABLE:
            raise RuntimeError("beautifulsoup4 required for URL conversion")
        
        # Fetch URL
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
        }
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        
        # Parse HTML
        soup = BeautifulSoup(response.content, 'html.parser')
        
        # Remove non-content elements
        for element in soup(['script', 'style', 'nav', 'footer', 'header', 'aside', 'ad']):
            element.decompose()
        
        # Try to find main content
        main_content = (
            soup.find('article') or
            soup.find('main') or
            soup.find(class_=re.compile(r'(content|article|post)', re.I)) or
            soup.body
        )
        
        if main_content:
            text = main_content.get_text(separator='\n')
        else:
            text = soup.get_text(separator='\n')
        
        # Clean up
        lines = [line.strip() for line in text.split('\n')]
        text = '\n'.join(line for line in lines if line)
        
        # Get title
        title_tag = soup.find('title')
        title = title_tag.string if title_tag else url
        
        info = DocumentInfo(
            format=DocumentFormat.URL,
            page_count=1,
            word_count=len(text.split()),
            has_images=len(soup.find_all('img')) > 0,
            has_tables=len(soup.find_all('table')) > 0,
            language=soup.html.get('lang') if soup.html else None,
            title=title,
            author=None,
            is_scanned=False
        )
        
        return text, info


# Singleton instances
document_converter = DocumentConverter()
url_converter = URLConverter()
